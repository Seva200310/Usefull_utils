# -*- coding: utf-8 -*-
import vk_api
import requests
import fpdf
from fpdf import FPDF
import docx
from datetime import datetime
import os
import vk

# if you encounter a "year is out of range" error the timestamp
# may be in milliseconds, try `ts /= 1000` in that case

fpdf.SYSTEM_TTFONTS = '/path/to/system/fonts'
""" Пример получения всех постов со стены """
login, password = '', ''
token = ""
vk_session= vk_api.VkApi(token=token)
vk=vk_session
api = vk.get_api()
try:
    vk_session.auth(token_only=True)
except vk_api.AuthError as error_msg:
    print(error_msg)


tools = vk_api.VkTools(vk_session)
def photoURls(post):
    #print(post)
    attachments=post['attachments']#[0]['photo']['sizes'][-1]['url']
    photos=[]
    URLs=[]    
    for attachment in attachments:
        if attachment['type']=='photo':
            photos.append(attachment['photo'])
    for photo in photos:
        URLs.append(photo['sizes'][-1]['url'])
    return URLs
def installPhoto(url,name):
    img_data = requests.get(url).content
    with open(f'{name}.jpg', 'wb') as handler:
        handler.write(img_data)
def makepostPagedata(wall_list):
    postPageData=[]
    #print(wall_list)
    for post in wall_list:
        #вытаскиваем текст
        text=post['text']
        ts=int(post['date'])
        datatime=datetime.utcfromtimestamp(ts).strftime('%Y %m %d %H:%M:%S')
        post_hash=post['hash']
        #вытаскиваем юрл фотok
        urls=photoURls(post)
        if text!='' or urls!=[]:
            #сскачаем фотки
            i=0
            file_names=[]
            for url in urls:
                name=post_hash+'-'+str(i)
                installPhoto(url,name)
                file_names.append(name)
                ++i
            postPage={'datatime':datatime,'text':text,'photos':file_names}#сохраняем имена фоток и текст в словари
            postPageData.append(postPage)
    return postPageData
        
    
def makeFirstPage(ID):
    pass
def make_pdf(postPageData):
    pass
def get_all(owner_id):


    """ VkTools.get_all позволяет получить все объекты со всех страниц.
        Соответственно get_all используется только если метод принимает
        параметры: count и offset.
        Например может использоваться для получения всех постов стены,
        всех диалогов, всех сообщений, etc.
        При использовании get_all сокращается количество запросов к API
        за счет метода execute в 25 раз.
        Например за раз со стены можно получить 100 * 25 = 2500, где
        100 - максимальное количество постов, которое можно получить за один
        запрос (обычно написано на странице с описанием метода)
    """

    wall = tools.get_all('wall.get', 100, {'owner_id': owner_id })

    print('Posts count:', wall['count'])

    if wall['count']:
        print('First post:', wall['items'][0], '\n')
        #url=photoURls(wall['items'][0])[0]
        #installPhoto(url,'photo.jpg')
    if wall['count'] > 1:
        print('Last post:', wall['items'][-1]['text'])
    return wall['items']

'''def makePDF(postPageData):
    pdf = FPDF()
    pdf.add_font('DejaVu', '', 'DejaVuSansCondensed.ttf', uni=True)
    pdf.set_font('DejaVu', '', 10) 
    pdf.add_page()
    line_no = 1
    for post in postPageData:
        pdf.cell(0, 10, txt=post['text'].format(line_no), ln=1)
        line_no += 1
    pdf.output("multipage_simple.pdf")'''
def makeDocx(ID):
    user_get = api.users.get(user_ids=(ID),fields='photo_200')
    user_get = user_get[0]
    print(user_get)
    first_name = user_get['first_name'] #Имя пользователя
    last_name = user_get['last_name'] #Фамилия
    photo=user_get['photo_200']
    installPhoto(photo,'ava')
    
    #profile_photo=
    doc = docx.Document()
    doc.add_picture('ava.jpg')
    os.remove('ava.jpg')
    head=doc.add_heading(f'id:{ID}\n имя:{first_name}\n фамилия:{last_name}').bold=True
    #head.alignment=1
    postPageData=makepostPagedata(get_all(ID))
    doc.paragraphs[1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
    for post in postPageData:
    # добавляем первый параграф
        #if post['text']!='' or post['photos']!=[]:
        doc.add_heading(post['datatime'], 2).bold=True
        doc.add_paragraph(post['text'])
        photos=post['photos']
        for photo in photos:
            try:
                paragraph=doc.add_picture(f'{photo}.jpg', width = docx.shared.Cm(10))
                #paragraph=paragraph+1
                os.remove(f'{photo}.jpg')
            except:
                pass
        doc.add_page_break() 
    doc.save(f'{ID}.docx')
    
res=makeDocx(345297171)
print('готово')
#print(res)

 

